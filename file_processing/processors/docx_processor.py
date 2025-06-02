from io import BytesIO
import msoffcrypto
from docx import Document
from file_processing.errors import FileProcessingFailedError, FileCorruptionError
from file_processing.file_processor_strategy import FileProcessorStrategy

import logging
logger = logging.getLogger(__name__)

class DocxFileProcessor(FileProcessorStrategy):
    """
    Processor for handling .docx files, extracting metadata and saving modifications.

    Attributes:
        metadata (dict): Contains metadata fields such as 'text', 'author', 'last_modified_by',
                         and 'has_password' if the file is opened.
    """

    def __init__(self, file_path: str, open_file: bool = True) -> None:
        """
        Initializes the DocxFileProcessor with the specified file path.

        Args:
            file_path (str): Path to the .docx file to process.
            open_file (bool): Indicates whether to open and process the file immediately.

        Sets:
            metadata (dict): Populated with a message if `open_file` is False, or with default metadata otherwise.
        """
        super().__init__(file_path, open_file)
        self.metadata = {'message': 'File was not opened'} if not open_file else self._default_metadata()

    def _default_metadata(self) -> dict:
        """
        Provides default metadata structure for a .docx file.

        Returns:
            dict: Default metadata fields with None or False values.
        """
        return {
            'text': None,
            'author': None,
            'last_modified_by': None,
            'has_password': False
        }

    def process(self) -> None:
        """
        Extracts metadata from the .docx file, including text content, author, and last modified information.
        Checks if the file is password-protected.

        Raises:
            FileCorruptionError: If the file is corrupted.
            FileProcessingFailedError: If an error occurs during file processing.
        """
        logger.info(f"Starting processing of DOCX file '{self.file_path}'.")

        if not self.open_file:
            logger.debug(f"DOCX file '{self.file_path}' was not opened (open_file=False).")
            return

        with open(self.file_path, 'rb') as f:
            file_content = BytesIO(f.read())

        try:
            office_file = msoffcrypto.OfficeFile(file_content)
            if office_file.is_encrypted():
                self.metadata["has_password"] = True
                logger.debug(f"DOCX file '{self.file_path}' is encrypted; skipping text extraction.")
                return
        except Exception as e:
            logger.error(f"Failed to process DOCX file '{self.file_path}': {e}")
            raise FileCorruptionError(f"File is corrupted: {self.file_path}") from e

        try:
            file_content.seek(0)  # Reset the position to the start
            doc = Document(file_content)
            self.metadata.update({'text': self.extract_text_from_docx(doc)})
            self.metadata.update({'author': doc.core_properties.author})
            self.metadata.update({'last_modified_by': doc.core_properties.last_modified_by})
            logger.debug(f"Extracted metadata from '{self.file_path}': Author='{doc.core_properties.author}', LastModifiedBy='{doc.core_properties.last_modified_by}'")
            logger.info(f"Successfully processed DOCX file '{self.file_path}'.")
        except Exception as e:
            logger.error(f"Failed to process DOCX file '{self.file_path}': {e}")
            raise FileProcessingFailedError(
                f"Error encountered while processing {self.file_path}: {e}"
            )

    def save(self, output_path: str = None) -> None:
        """
        Saves updated metadata to the .docx file at the specified output path.

        Args:
            output_path (str): Path to save the file with updated metadata. If None, overwrites the original file.

        Raises:
            FileProcessingFailedError: If an error occurs while saving the .docx file.
        """
        save_path = output_path or self.file_path
        logger.info(f"Saving DOCX file '{self.file_path}' to '{save_path}'.")

        try:
            doc = Document(self.file_path)

            # Update the core properties (metadata)
            cp = doc.core_properties
            cp.author = self.metadata.get('author', cp.author)
            cp.last_modified_by = self.metadata.get('last_modified_by', cp.last_modified_by)

            doc.save(save_path)
            logger.info(f"DOCX file '{self.file_path}' saved successfully to '{save_path}'.")
        except Exception as e:
            logger.error(f"Failed to save DOCX file '{self.file_path}' to '{save_path}': {e}")
            raise FileProcessingFailedError(
                f"Error encountered while saving to {save_path}: {e}"
            )

    def extract_text_from_docx(self, doc: Document) -> str:
        """
        Extracts and returns the full text content from a .docx Document object.

        Args:
            doc (Document): A `Document` instance representing the .docx file.

        Returns:
            str: Extracted text content from the .docx file.

        Raises:
            FileProcessingFailedError: If an error occurs while reading the document.
        """
        try:
            full_text = [para.text for para in doc.paragraphs]
            return '\n'.join(full_text)
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX file '{self.file_path}': {e}")
            raise FileProcessingFailedError(
                f"Error encountered while opening or processing {self.file_path}: {e}"
            )
