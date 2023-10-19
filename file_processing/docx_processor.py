from file_processor_strategy import FileProcessorStrategy
from docx import Document
from zipfile import BadZipFile
from docx.oxml import OxmlElement
from errors import FileProcessingFailedError, FileCorruptionError
from io import BytesIO

class DocxFileProcessor(FileProcessorStrategy):
    def __init__(self, file_path: str) -> None:
        super().__init__(file_path)
        self.metadata = self._default_metadata()


    def _default_metadata(self) -> dict:
        return {
            'text': None,
            'author': None,
            'last_modified_by': None,
            'has_password': False
        }


    def process(self) -> None:

        try:
            with open(self.file_path, 'rb') as f:
                file_content = BytesIO(f.read())
        except OSError as e:
            self.metadata["has_password"] = True
            return

        try:
            doc = Document(file_content)
            self.metadata.update({'text': self.extract_text_from_docx(doc)})
            self.metadata.update({'author': doc.core_properties.author})
            self.metadata.update({'last_modified_by': doc.core_properties.last_modified_by})
        except BadZipFile as e:
            raise FileCorruptionError(f"Error encountered while processing {self.file_path}: {e}")

        # Other core properties to include: https://python-docx.readthedocs.io/en/latest/api/document.html#coreproperties-objects
        # keywords, language, subject, version


    def save(self, output_path: str = None) -> None:
        doc = Document(self.file_path)

        # Update the core properties (metadata)
        cp = doc.core_properties
        cp.author = self.metadata.get('author', cp.author)
        cp.last_modified_by = self.metadata.get('last_modified_by', cp.last_modified_by)
        
        save_path = output_path or self.file_path
        doc.save(save_path)


    @staticmethod
    def extract_text_from_docx(doc: Document) -> str:
        try:
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            return '\n'.join(full_text)
        except Exception as e:
            print(f"Error encountered while opening or processing {self.file_path}: {e}")
            return None